# Databricks notebook source
# MAGIC %md
# MAGIC # 08 - Unstructured Data: Document Processing on FSxN
# MAGIC
# MAGIC Extract text and metadata from PDF/DOCX documents stored on FSx for NetApp ONTAP.
# MAGIC Build a searchable document catalog as a Delta table.
# MAGIC
# MAGIC ## Use Cases
# MAGIC - Document metadata cataloging (page count, author, creation date)
# MAGIC - Text extraction for full-text search indexing
# MAGIC - RAG (Retrieval-Augmented Generation) pipeline preparation
# MAGIC - Compliance document inventory
# MAGIC
# MAGIC ## Prerequisites
# MAGIC - ML Runtime cluster (Python libraries available)
# MAGIC - External Location configured (from notebook 01)
# MAGIC - Sample PDF/DOCX documents uploaded to FSxN via NFS/SMB
# MAGIC - Install: `%pip install pypdf python-docx`

# COMMAND ----------

# MAGIC %pip install pypdf python-docx

# COMMAND ----------

# MAGIC %md
# MAGIC ## Configuration

# COMMAND ----------

S3_ACCESS_POINT_ALIAS = "<your-s3ap-alias>"
DOCUMENTS_PATH = f"s3://{S3_ACCESS_POINT_ALIAS}/media/documents/"
CATALOG_OUTPUT_PATH = f"s3://{S3_ACCESS_POINT_ALIAS}/silver/document_catalog/"
TEXT_OUTPUT_PATH = f"s3://{S3_ACCESS_POINT_ALIAS}/silver/document_text/"

# COMMAND ----------

# MAGIC %md
# MAGIC ## Step 1: Read Documents as Binary Files

# COMMAND ----------

# Read PDF and DOCX files
docs_df = spark.read.format("binaryFile") \
    .option("pathGlobFilter", "*.{pdf,docx}") \
    .option("recursiveFileLookup", "true") \
    .load(DOCUMENTS_PATH)

print(f"Found {docs_df.count()} documents")
docs_df.select("path", "length", "modificationTime").show(10, truncate=False)

# COMMAND ----------

# MAGIC %md
# MAGIC ## Step 2: Extract PDF Metadata and Text

# COMMAND ----------

from pyspark.sql.functions import udf, col, current_timestamp, when, lit
from pyspark.sql.types import StructType, StructField, StringType, IntegerType, ArrayType
from io import BytesIO

# Schema for document metadata
doc_metadata_schema = StructType([
    StructField("page_count", IntegerType()),
    StructField("author", StringType()),
    StructField("title", StringType()),
    StructField("subject", StringType()),
    StructField("creator", StringType()),
    StructField("creation_date", StringType()),
    StructField("text_preview", StringType()),  # First 500 chars
    StructField("total_chars", IntegerType()),
    StructField("doc_type", StringType()),
    StructField("error", StringType()),
])

@udf(returnType=doc_metadata_schema)
def extract_document_metadata(content, path):
    """Extract metadata and text preview from PDF or DOCX."""
    try:
        if path.lower().endswith('.pdf'):
            return _extract_pdf(content)
        elif path.lower().endswith('.docx'):
            return _extract_docx(content)
        else:
            return (0, None, None, None, None, None, None, 0, "unknown", "Unsupported format")
    except Exception as e:
        return (0, None, None, None, None, None, None, 0, "error", str(e)[:200])

def _extract_pdf(content):
    """Extract from PDF using pypdf."""
    from pypdf import PdfReader
    reader = PdfReader(BytesIO(content))
    info = reader.metadata

    # Extract text from all pages
    full_text = ""
    for page in reader.pages:
        text = page.extract_text()
        if text:
            full_text += text + "\n"

    return (
        len(reader.pages),
        str(info.author) if info and info.author else None,
        str(info.title) if info and info.title else None,
        str(info.subject) if info and info.subject else None,
        str(info.creator) if info and info.creator else None,
        str(info.creation_date) if info and info.creation_date else None,
        full_text[:500] if full_text else None,
        len(full_text),
        "pdf",
        None
    )

def _extract_docx(content):
    """Extract from DOCX using python-docx."""
    from docx import Document
    doc = Document(BytesIO(content))

    # Extract text from paragraphs
    full_text = "\n".join([para.text for para in doc.paragraphs if para.text])

    # Extract core properties
    props = doc.core_properties
    return (
        len(doc.paragraphs),  # Approximate "pages" as paragraph count
        props.author if props.author else None,
        props.title if props.title else None,
        props.subject if props.subject else None,
        None,  # creator not standard in docx
        str(props.created) if props.created else None,
        full_text[:500] if full_text else None,
        len(full_text),
        "docx",
        None
    )

# COMMAND ----------

# Apply extraction
catalog_df = docs_df.withColumn(
    "metadata", extract_document_metadata(col("content"), col("path"))
).select(
    col("path"),
    col("length").alias("file_size_bytes"),
    col("modificationTime").alias("last_modified"),
    col("metadata.page_count"),
    col("metadata.author"),
    col("metadata.title"),
    col("metadata.subject"),
    col("metadata.creator"),
    col("metadata.creation_date"),
    col("metadata.text_preview"),
    col("metadata.total_chars"),
    col("metadata.doc_type"),
    col("metadata.error"),
    current_timestamp().alias("processed_at")
)

# Show results
catalog_df.select("path", "doc_type", "page_count", "total_chars", "title").show(10, truncate=False)

# COMMAND ----------

# MAGIC %md
# MAGIC ## Step 3: Save Document Catalog as Delta Table

# COMMAND ----------

# Write catalog (without full text, just metadata + preview)
catalog_df.drop("text_preview").write \
    .format("delta") \
    .mode("overwrite") \
    .save(CATALOG_OUTPUT_PATH)

spark.sql(f"""
CREATE TABLE IF NOT EXISTS fsxn_lakehouse.silver.document_catalog
USING DELTA
LOCATION '{CATALOG_OUTPUT_PATH}'
COMMENT 'Document metadata catalog from FSxN unstructured files'
""")

print(f"✅ Document catalog saved: {CATALOG_OUTPUT_PATH}")

# COMMAND ----------

# MAGIC %md
# MAGIC ## Step 4: Extract Full Text for RAG Pipeline

# COMMAND ----------

@udf(returnType=StringType())
def extract_full_text(content, path):
    """Extract complete text from document."""
    try:
        if path.lower().endswith('.pdf'):
            from pypdf import PdfReader
            reader = PdfReader(BytesIO(content))
            return "\n".join([
                page.extract_text() or "" for page in reader.pages
            ])
        elif path.lower().endswith('.docx'):
            from docx import Document
            doc = Document(BytesIO(content))
            return "\n".join([para.text for para in doc.paragraphs if para.text])
        return ""
    except Exception as e:
        return f"ERROR: {str(e)[:100]}"

# Extract full text
text_df = docs_df.select(
    col("path"),
    extract_full_text(col("content"), col("path")).alias("full_text"),
    current_timestamp().alias("extracted_at")
).filter(~col("full_text").startswith("ERROR"))

# Save as Delta table (for RAG indexing)
text_df.write \
    .format("delta") \
    .mode("overwrite") \
    .save(TEXT_OUTPUT_PATH)

spark.sql(f"""
CREATE TABLE IF NOT EXISTS fsxn_lakehouse.silver.document_text
USING DELTA
LOCATION '{TEXT_OUTPUT_PATH}'
COMMENT 'Full text extracted from FSxN documents (for RAG/search)'
""")

print(f"✅ Document text saved: {TEXT_OUTPUT_PATH}")

# COMMAND ----------

# MAGIC %md
# MAGIC ## Step 5: Text Chunking for RAG

# COMMAND ----------

from pyspark.sql.functions import explode, split, length, monotonically_increasing_id

# Simple chunking by paragraphs (for RAG embedding)
CHUNK_SIZE = 1000  # characters per chunk

@udf(returnType=ArrayType(StringType()))
def chunk_text(text, chunk_size=1000):
    """Split text into chunks for embedding."""
    if not text:
        return []
    chunks = []
    for i in range(0, len(text), chunk_size):
        chunk = text[i:i + chunk_size].strip()
        if chunk:
            chunks.append(chunk)
    return chunks

chunks_df = text_df.withColumn("chunks", chunk_text(col("full_text"))) \
    .select("path", explode("chunks").alias("chunk")) \
    .withColumn("chunk_id", monotonically_increasing_id()) \
    .withColumn("chunk_length", length("chunk"))

print(f"Total chunks: {chunks_df.count()}")
chunks_df.show(5, truncate=80)

# COMMAND ----------

# MAGIC %md
# MAGIC ## Summary
# MAGIC
# MAGIC | Output | Location | Format |
# MAGIC |--------|----------|--------|
# MAGIC | Document Catalog | `silver/document_catalog/` | Delta (metadata only) |
# MAGIC | Full Text | `silver/document_text/` | Delta (complete text) |
# MAGIC | Text Chunks | (in-memory for embedding) | DataFrame |
# MAGIC
# MAGIC ## ONTAP Value for Document Processing
# MAGIC
# MAGIC | Feature | Benefit |
# MAGIC |---------|---------|
# MAGIC | SnapLock | WORM protection for compliance documents |
# MAGIC | Snapshot | Preserve document state before batch processing |
# MAGIC | Multi-protocol | SMB (Windows users upload) + S3 AP (Databricks processes) |
# MAGIC | Deduplication | Multiple versions of same document share blocks |
# MAGIC | FlexClone | Isolated copy for RAG experimentation |
# MAGIC
# MAGIC ## Next Steps
# MAGIC - Run `09_ml_embeddings_on_fsxn.py` to generate embeddings from chunks
# MAGIC - Connect to Bedrock Knowledge Base for RAG queries
# MAGIC - Build full-text search index with OpenSearch
